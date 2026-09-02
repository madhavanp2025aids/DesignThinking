"""
HYDAC Spec-to-3D Generator — DOCX Parser
Atomic, testable module: one file type only.
Uses python-docx for table + paragraph text extraction.
"""

from typing import Optional


def parse_docx(file_path: str, filename: str) -> dict:
    """
    Parse a DOCX file, extracting text and tables in document order.

    Returns:
        {
            "raw_text": str,
            "tables": [{"table_index": int, "source_location": str, "headers": [...], "rows": [[...]]}],
            "paragraphs": [{"index": int, "text": str, "source_location": str}]
        }
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX parsing. Install: pip install python-docx")

    doc = Document(file_path)

    result = {
        "raw_text": "",
        "tables": [],
        "paragraphs": [],
    }

    table_index = 0
    para_index = 0
    element_index = 0

    # Extract headers and footers from all sections
    for sec_idx, section in enumerate(doc.sections, start=1):
        if section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                txt = p.text.strip()
                if txt:
                    result["paragraphs"].append({
                        "index": para_index,
                        "text": txt,
                        "source_location": f"Header (Section {sec_idx})",
                    })
                    result["raw_text"] += f"[Header]: {txt}\n"
                    para_index += 1

        if section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                txt = p.text.strip()
                if txt:
                    result["paragraphs"].append({
                        "index": para_index,
                        "text": txt,
                        "source_location": f"Footer (Section {sec_idx})",
                    })
                    result["raw_text"] += f"[Footer]: {txt}\n"
                    para_index += 1

    # Iterate document body in order (tables and paragraphs interleaved)
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "tbl":
            # This is a table element — find the matching Table object
            for table in doc.tables:
                if table._element is element:
                    table_data = _extract_table(table, table_index, element_index)
                    if table_data:
                        result["tables"].append(table_data)
                        table_index += 1
                    break
            element_index += 1

        elif tag == "p":
            # This is a paragraph element — find the matching Paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if text:
                        para_data = {
                            "index": para_index,
                            "text": text,
                            "source_location": f"Paragraph {para_index + 1}",
                        }
                        result["paragraphs"].append(para_data)
                        result["raw_text"] += text + "\n"
                        para_index += 1
                    break
            element_index += 1

    return result


def _extract_table(table, table_index: int, element_index: int) -> Optional[dict]:
    """Extract a single table from the document."""
    rows_data = []

    for row in table.rows:
        row_cells = []
        seen_texts = set()
        for cell in row.cells:
            # Handle merged cells: python-docx repeats merged cell references
            cell_text = cell.text.strip()
            cell_id = id(cell._element)
            if cell_id in seen_texts:
                row_cells.append("")  # Merged cell — don't duplicate
            else:
                seen_texts.add(cell_id)
                row_cells.append(cell_text)
        rows_data.append(row_cells)

    if len(rows_data) < 2:
        return None

    # First row as headers
    headers = rows_data[0]
    data_rows = rows_data[1:]

    return {
        "table_index": table_index,
        "source_location": f"Table {table_index + 1} (element {element_index})",
        "headers": headers,
        "rows": data_rows,
    }
